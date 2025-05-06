import glob
import PyPDF2
import docx
import chardet
from spire.doc import *
import os
import logging
'''
pip install PyPDF2
pip install python-docx
pip install Spire.Doc
pip install python-pptx
'''


# 获取指定文件夹中所有doc文件
def get_doc_files(root_dir):
    doc_files = []
    for root, dirs, files in os.walk(root_dir):
        for file in files:
            if file.endswith('.doc'):  # 查找 .doc 文件
                doc_files.append(os.path.join(root, file))
    return doc_files


def doc_to_text():
    doc_files = get_doc_files(r'data')
    for path in doc_files:
        if path.lower().endswith('doc'):
            doc = Document()
            # 加载.doc文件
            doc.LoadFromFile(path)
            # 提取全部文本
            text = doc.GetText()
            dir_name, file_name = os.path.split(path)
            file_name_without_ext, ext = os.path.splitext(file_name)
            new_file_name = file_name_without_ext + '.txt'
            path = os.path.join(dir_name, new_file_name)
            # 保存为txt文件
            with open(path, "w", encoding="utf-8") as f:
                f.write(text)
            doc.Close()
        else:
            pass


# 读取pdf文件
def read_pdf(path):
    text_list = []
    with open(path, "rb") as file:
        reader = PyPDF2.PdfReader(file)
        text = ""
        i = 0
        for page in reader.pages:
            if i <= 2:
                text += page.extract_text()
                i += 1
            else:
                text_list.append(text)
                i = 0
                text = ""
        text_list.append(text)

    return text_list


# 读取docx文件
def read_docx(file_path):
    text_list = []
    doc = docx.Document(file_path)
    text = ''
    i = 0
    for paragraph in doc.paragraphs:
        if i <= 5:
            text += paragraph.text
            i += 1
        else:
            text_list.append(text)
            i = 0
            text = ''
    text_list.append(text)
    return text_list


def detect_encoding(file_path):
    with open(file_path, 'rb') as f:
        result = chardet.detect(f.read())
    return result['encoding']


# 读取txt文本
def read_txt(file_path):
    text_list = []
    encoding = detect_encoding(file_path)
    with open(file_path, 'r', encoding=encoding) as f:
        for line in f.readlines():
            text_list.append(line)
    return text_list


def read_data(path_list):
    logging.basicConfig(filename='log.log', level=logging.INFO,
                        format='%(asctime)s - %(levelname)s - %(message)s')
    doc_to_text()
    text_list = []
    for path in path_list:
        try:
            if path.lower().endswith('pdf'):
                text = read_pdf(path)
                text_list += text
            if path.lower().endswith('docx'):
                text = read_docx(path)
                text_list += text
            if path.lower().endswith('txt'):
                text = read_txt(path)
                text_list += text
        except Exception as e:
            logging.error(e)
            continue

    return text_list


if __name__ == '__main__':
    doc_to_text()
    path_list = glob.glob("data/*.*")
    text_list = read_data(path_list)
    print(text_list)
